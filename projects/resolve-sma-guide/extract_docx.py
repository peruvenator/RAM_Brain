"""Extract ReSolve SMA Guide .docx to markdown with frontmatter."""

from docx import Document
from pathlib import Path


def runs_to_md(runs):
    """Convert a list of runs to markdown with bold/italic."""
    parts = []
    for r in runs:
        text = r.text
        if not text:
            continue
        if r.bold and r.italic:
            text = f"***{text}***"
        elif r.bold:
            text = f"**{text}**"
        elif r.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


def table_to_md(table):
    """Convert a Word table to markdown table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # Check if it's a single-cell "text block" table (like the email templates)
    if len(rows) == 1 and len(rows[0]) == 1:
        content = table.rows[0].cells[0].text.strip()
        return f"> {content}\n"

    # Build markdown table
    lines = []
    # Header
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    # Data rows
    for row in rows[1:]:
        # Pad row if needed
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines) + "\n"


def extract():
    here = Path(__file__).parent
    doc = Document(here / "ReSolve SMA Guide - Client Facing v1.0.docx")

    # Build an ordered list of content elements (paragraphs and tables)
    # by walking the document body XML
    from docx.oxml.ns import qn

    body = doc.element.body
    para_idx = 0
    table_idx = 0
    elements = []

    for child in body:
        if child.tag == qn("w:p"):
            elements.append(("para", doc.paragraphs[para_idx]))
            para_idx += 1
        elif child.tag == qn("w:tbl"):
            elements.append(("table", doc.tables[table_idx]))
            table_idx += 1

    # Frontmatter
    md_lines = [
        "---",
        'title: "ReSolve SMA Guide"',
        'subtitle: "Understanding Your ReSolve SMA: Structure, Strategies, and Operations"',
        'date: "2026-04-01"',
        'cover-date: "March 2026"',
        'report-type: "CLIENT GUIDE"',
        "---",
        "",
    ]

    # Skip cover page paragraphs (first 5 lines are title/subtitle/version)
    skip_indices = {0, 1, 2, 3, 4}  # cover material
    para_counter = 0

    for elem_type, elem in elements:
        if elem_type == "para":
            p = elem
            idx = para_counter
            para_counter += 1

            if idx in skip_indices:
                continue

            text = p.text.strip()
            if not text:
                md_lines.append("")
                continue

            style = p.style.name if p.style else ""

            if style == "Heading 1":
                md_lines.append(f"## {text}")
                md_lines.append("")
            elif style == "Heading 2":
                md_lines.append(f"### {text}")
                md_lines.append("")
            elif style == "Heading 3":
                md_lines.append(f"#### {text}")
                md_lines.append("")
            elif style == "List Paragraph":
                md_text = runs_to_md(p.runs)
                md_lines.append(f"- {md_text}")
            else:
                md_text = runs_to_md(p.runs)
                md_lines.append(md_text)
                md_lines.append("")

        elif elem_type == "table":
            md_lines.append("")
            md_lines.append(table_to_md(elem))
            md_lines.append("")

    output = "\n".join(md_lines)
    # Clean up excessive blank lines
    while "\n\n\n\n" in output:
        output = output.replace("\n\n\n\n", "\n\n\n")

    out_path = here / "draft.md"
    out_path.write_text(output, encoding="utf-8")
    print(f"Wrote {len(output)} chars to {out_path}")


if __name__ == "__main__":
    extract()
