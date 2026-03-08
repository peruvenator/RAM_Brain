"""
ReSolve AI Tool Security Assessment Document Generator
=======================================================
Generates Word documents matching the ReSolve compliance template format.

Usage:
    Called by Claude Code with tool data to produce .docx files.
    Can also be run standalone with a JSON config file:
        python generate_assessment.py config.json
"""

from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import json
import sys
from datetime import date

# ── Constants ────────────────────────────────────────────────────────────────

DARK_BLUE = "1B3A5C"
LIGHT_GRAY = "F5F5F5"
WHITE = "FFFFFF"
SUBTITLE_BLUE = "B0C4DE"
TEXT_COLOR = "333333"

LINK_BLUE = "0563C1"
REVIEW_TAG = " subject to review by appropriate personnel"

FONT_NAME = "Arial"
HEADER_SIZE = Pt(11)       # Section headers (139700 EMU)
BODY_SIZE = Pt(9.5)        # Body text (120650 EMU)
TITLE_SIZE = Pt(14)        # Banner title (177800 EMU)
SUBTITLE_SIZE = Pt(11)     # Banner subtitle

# Page setup (Letter, narrow margins)
PAGE_WIDTH = Emu(7772400)
PAGE_HEIGHT = Emu(10058400)
MARGIN_TOP = Emu(457200)
MARGIN_BOTTOM = Emu(457200)
MARGIN_LEFT = Emu(571500)
MARGIN_RIGHT = Emu(571500)


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_width(cell, width_dxa):
    """Set explicit cell width in twips (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>'
    )
    tcPr.append(tcW)


def add_run(paragraph, text, bold=None, size=BODY_SIZE, color_hex=TEXT_COLOR, font_name=FONT_NAME):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color_hex)
    return run


def add_review_tag(paragraph, size=BODY_SIZE, font_name=FONT_NAME):
    """Append ' subject to review by appropriate personnel' in standard body text."""
    run = paragraph.add_run(REVIEW_TAG)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = RGBColor.from_string(TEXT_COLOR)
    return run


def add_hyperlink(paragraph, url, display_text, size=BODY_SIZE, color_hex="0563C1", font_name=FONT_NAME):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'</w:hyperlink>'
    )

    run_elem = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr>'
        f'    <w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'    <w:sz w:val="{int(size.pt * 2)}"/>'
        f'    <w:szCs w:val="{int(size.pt * 2)}"/>'
        f'    <w:color w:val="{color_hex}"/>'
        f'    <w:u w:val="single"/>'
        f'    <w:rStyle w:val="Hyperlink"/>'
        f'  </w:rPr>'
        f'  <w:t xml:space="preserve">{display_text}</w:t>'
        f'</w:r>'
    )

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def format_header_row(row, col_texts, col_count=2):
    """Format a table header row with dark blue background and white text."""
    for i, cell in enumerate(row.cells):
        set_cell_shading(cell, DARK_BLUE)
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        add_run(para, col_texts[i], bold=True, color_hex=WHITE)


def format_data_row(row, col_data, row_index, has_hyperlink=False):
    """Format a data row with alternating shading. col_data is list of (text, bold) tuples."""
    is_shaded = (row_index % 2 == 0)  # Even data rows (0-indexed) get shading
    for i, cell in enumerate(row.cells):
        if is_shaded:
            set_cell_shading(cell, LIGHT_GRAY)
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        text, bold = col_data[i]
        if has_hyperlink and i == 1 and text.startswith("http"):
            add_hyperlink(para, text, text)
        else:
            add_run(para, text, bold=bold)


def add_section_header(doc, text):
    """Add a section header paragraph (bold, dark blue, 11pt)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    add_run(para, text, bold=True, size=HEADER_SIZE, color_hex=DARK_BLUE)
    return para


def create_two_col_table(doc, headers, rows, col_widths=(2800, 7640), has_hyperlinks=False):
    """Create a standard 2-column table with header row and alternating shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False

    # Header row
    format_header_row(table.rows[0], headers)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        format_data_row(row, row_data, ri, has_hyperlink=has_hyperlinks)

    # Set column widths
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if ci < len(col_widths):
                set_cell_width(cell, col_widths[ci])

    # Remove default table borders and set custom thin borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    return table


# ── Main Generator ───────────────────────────────────────────────────────────

def generate_assessment(data, output_path):
    """
    Generate a ReSolve AI Tool Security Assessment Word document.

    Parameters:
        data (dict): Tool assessment data with the following structure:
            {
                "tool_name_short": "ChatGPT Business",
                "vendor_short": "OpenAI",
                "banner_title": "ChatGPT Business by OpenAI",
                "overview": {
                    "tool_name": "ChatGPT Business (Web / Desktop / Mobile)",
                    "vendor": "OpenAI  —  San Francisco, CA",
                    "plan_tier": "Business ($25/seat/mo annual; $30/seat/mo monthly)",
                    "internal_owner": "Rodrigo Gordillo, President",
                    "assessment_date": "February 2026",
                    "policy_reference": "ReSolve Compliance Manual §VIII.G; ReSolve Artificial Intelligence Policy (Orion)"
                },
                "tool_description": "ChatGPT Business is a team-focused AI workspace...",
                "use_cases": [
                    {"text": "Drafting investment research memos...", "review": true},
                    "Analyzing financial data...",
                    {"text": "Composing client communications...", "review": true}
                ],
                "data_protection": [
                    {"control": "Model Training", "details": "Disabled by default..."},
                    {"control": "Encryption — In Transit", "details": "TLS 1.2+..."},
                    ...
                ],
                "certifications": [
                    {"certification": "SOC 2 Type II", "status": "Compliant..."},
                    ...
                ],
                "security_links": [
                    {"resource": "Trust Portal", "url": "https://trust.openai.com"},
                    ...
                ],
                "approval_roles": [
                    "Tool Sponsor",
                    "Cybersecurity",
                    "Compliance"
                ]
            }
        output_path (str): Path for the output .docx file.
    """
    doc = Document()

    # ── Page Setup ───────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # ── Table 0: Title Banner ────────────────────────────────────────────
    banner = doc.add_table(rows=1, cols=1)
    banner.autofit = False
    cell = banner.rows[0].cells[0]
    set_cell_shading(cell, DARK_BLUE)

    # Title line
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    add_run(para, "AI TOOL SECURITY ASSESSMENT", bold=True, size=TITLE_SIZE, color_hex=WHITE)

    # Subtitle line
    para2 = cell.add_paragraph()
    para2.paragraph_format.space_before = Pt(0)
    para2.paragraph_format.space_after = Pt(6)
    add_run(para2, data["banner_title"], size=SUBTITLE_SIZE, color_hex=SUBTITLE_BLUE)

    # Remove banner borders
    tbl = banner._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # ── Table 1: Overview ────────────────────────────────────────────────
    doc.add_paragraph()  # Spacer
    overview = data["overview"]
    overview_rows = [
        [("Tool Name", True), (overview["tool_name"], None)],
        [("Vendor", True), (overview["vendor"], None)],
        [("Plan / Tier", True), (overview["plan_tier"], None)],
        [("Internal Owner", True), (overview["internal_owner"], None)],
        [("Assessment Date", True), (overview["assessment_date"], None)],
        [("Policy Reference", True), (overview["policy_reference"], None)],
    ]
    create_two_col_table(doc, ["Field", "Details"], overview_rows)

    # ── Tool Description ─────────────────────────────────────────────────
    add_section_header(doc, "Tool Description")
    desc_para = doc.add_paragraph()
    desc_para.paragraph_format.space_before = Pt(0)
    desc_para.paragraph_format.space_after = Pt(6)
    add_run(desc_para, data["tool_description"])

    # ── Use Cases ────────────────────────────────────────────────────────
    # Each use case can be a plain string or a dict:
    #   {"text": "Drafting memos...", "review": true}
    # When "review" is true, appends blue underlined review tag.
    add_section_header(doc, "Use Cases")
    for case in data["use_cases"]:
        para = doc.add_paragraph(style="List Bullet")
        if para.runs:
            para.runs[0].text = ""
        if isinstance(case, dict):
            add_run(para, case["text"])
            if case.get("review"):
                add_review_tag(para)
        else:
            add_run(para, case)

    # ── Data Protection & Privacy ────────────────────────────────────────
    add_section_header(doc, "Data Protection & Privacy")
    dp_rows = [
        [(item["control"], True), (item["details"], None)]
        for item in data["data_protection"]
    ]
    create_two_col_table(doc, ["Control", "Details"], dp_rows)

    # ── Vendor Security & Compliance ─────────────────────────────────────
    add_section_header(doc, "Vendor Security & Compliance")
    cert_rows = [
        [(item["certification"], True), (item["status"], None)]
        for item in data["certifications"]
    ]
    create_two_col_table(doc, ["Certification", "Status"], cert_rows)

    # ── Vendor Security Documentation ────────────────────────────────────
    add_section_header(doc, "Vendor Security Documentation")

    # Build links table with hyperlinks
    links_table = doc.add_table(rows=1 + len(data["security_links"]), cols=2)
    links_table.autofit = False

    # Header row
    format_header_row(links_table.rows[0], ["Resource", "Link"])

    # Link rows
    for ri, link_data in enumerate(data["security_links"]):
        row = links_table.rows[ri + 1]
        is_shaded = (ri % 2 == 0)

        # Resource cell
        cell0 = row.cells[0]
        if is_shaded:
            set_cell_shading(cell0, LIGHT_GRAY)
        para0 = cell0.paragraphs[0]
        para0.paragraph_format.space_before = Pt(2)
        para0.paragraph_format.space_after = Pt(2)
        add_run(para0, link_data["resource"], bold=True)

        # Link cell with hyperlink
        cell1 = row.cells[1]
        if is_shaded:
            set_cell_shading(cell1, LIGHT_GRAY)
        para1 = cell1.paragraphs[0]
        para1.paragraph_format.space_before = Pt(2)
        para1.paragraph_format.space_after = Pt(2)
        add_hyperlink(para1, link_data["url"], link_data["url"])

        # Set column widths
        set_cell_width(cell0, 2800)
        set_cell_width(cell1, 7640)

    # Set header column widths too
    set_cell_width(links_table.rows[0].cells[0], 2800)
    set_cell_width(links_table.rows[0].cells[1], 7640)

    # Table borders
    tbl = links_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # ── Approval ─────────────────────────────────────────────────────────
    add_section_header(doc, "Approval")
    roles = data.get("approval_roles", ["Tool Sponsor", "Cybersecurity", "Compliance"])
    approval_table = doc.add_table(rows=1 + len(roles), cols=3)
    approval_table.autofit = False

    # Header
    format_header_row(approval_table.rows[0], ["Role", "Name / Signature", "Date"])

    # Rows
    today_str = date.today().strftime("%B %d, %Y")
    for ri, role in enumerate(roles):
        row = approval_table.rows[ri + 1]
        is_shaded = (ri % 2 == 0)
        for ci, cell in enumerate(row.cells):
            if is_shaded:
                set_cell_shading(cell, LIGHT_GRAY)
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            if ci == 0:
                add_run(para, role, bold=True)
            elif ci == 1 and role == "Tool Sponsor":
                add_run(para, "Rodrigo Gordillo")
            elif ci == 2 and role == "Tool Sponsor":
                add_run(para, today_str)

    # Set column widths for approval table
    for row in approval_table.rows:
        set_cell_width(row.cells[0], 2400)
        set_cell_width(row.cells[1], 5040)
        set_cell_width(row.cells[2], 3000)

    # Table borders
    tbl = approval_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(output_path)
    return output_path


# ── CLI Entry Point ──────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_assessment.py <config.json> [output.docx]")
        print("\nProvide a JSON file with tool assessment data.")
        sys.exit(1)

    config_path = sys.argv[1]
    with open(config_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Output path: either specified or auto-generated
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        safe_name = data.get("tool_name_short", "Tool").replace(" ", "_")
        output_path = f"ReSolve_AI_Security_Assessment_{safe_name}.docx"

    result = generate_assessment(data, output_path)
    print(f"Generated: {result}")
